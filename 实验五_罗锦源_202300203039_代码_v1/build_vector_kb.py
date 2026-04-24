import argparse
import json
import pickle
import os
import time
import shutil
import logging
from pathlib import Path

# 配置日志
logger = logging.getLogger("vector_kb")
from typing import List, Dict, Any, Tuple

from preprocess_docs import normalize_text

# 引入 LangChain 和 ChromaDB
from langchain_chroma import Chroma
from langchain_ollama import OllamaEmbeddings
from langchain_core.documents import Document
import chromadb

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}
QUERY_HINTS = {
    "周期": "periodization block periodized training cycle macrocycle mesocycle",
    "营养": "nutrition marathon hydration carbohydrate fueling diet protein",
    "饮食": "diet nutrition food eating meal plan calorie",
    "补给": "fueling hydration carbohydrate electrolyte intake",
    "过度训练": "overtraining overreaching excessive loading fatigue injury",
    "伤": "injury recovery rehabilitation pain therapist orthopedic",
    "痛": "pain injury recovery rehabilitation therapist",
    "康复": "recovery rehabilitation injury physical therapy",
    "高温": "heat stress hot environment hydration safety thermal",
    "力量": "strength training resistance program gym weightlifting",
    "耐力": "endurance training aerobic running cardiovascular",
    "马拉松": "marathon running training plan race long distance",
    "配速": "pace speed velocity timing splits",
    "训练": "training exercise workout program",
    "安全": "safety risk warning caution advice health emergency",
    "风险": "risk danger safety hazard warning",
    "禁忌": "contraindication warning caution safety avoid",
    "知识库": "knowledge base scope domain expert area content coverage",
    "领域": "knowledge base scope domain expert area content coverage",
    "文档": "book document source reference literature paper pdf",
    "参考": "reference source bibliography citations documentation",
    "原则": "principle methodology theory fundamental framework guidelines",
}
DEFAULT_TEST_QUESTIONS = [
    "如何定义并安排周期化训练以支持长期跑步表现提升？",
    "马拉松训练期如何进行营养与补给规划？",
    "过度训练常见风险信号有哪些，如何识别？",
    "高温环境下训练有哪些安全建议？",
    "力量训练与耐力训练如何在同一训练计划中平衡？",
]

OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
EMBEDDING_MODEL = "nomic-embed-text" # 可以根据实际安装的模型替换

def get_embeddings():
    return OllamaEmbeddings(model=EMBEDDING_MODEL, base_url=OLLAMA_BASE_URL)

def extract_pdf_pages(file_path: Path) -> list[tuple[int, str]]:
    try:
        import fitz
    except Exception:
        fitz = None
    if fitz is not None:
        pages = []
        with fitz.open(file_path) as doc:
            for idx, page in enumerate(doc, start=1):
                page_text = page.get_text("text") or ""
                pages.append((idx, page_text))
        return pages
    try:
        from pypdf import PdfReader
    except Exception as exc:
        raise RuntimeError("未安装可用的 PDF 文本抽取库，请安装 PyMuPDF 或 pypdf") from exc
    reader = PdfReader(str(file_path))
    pages = []
    for idx, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        pages.append((idx, page_text))
    return pages

def extract_docx_text(file_path: Path) -> str:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("处理 docx 需要安装 python-docx") from exc
    document = Document(str(file_path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text)

def load_pages(file_path: Path) -> list[tuple[int, str]]:
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return extract_pdf_pages(file_path)
    if ext in {".txt", ".md"}:
        return [(1, file_path.read_text(encoding="utf-8", errors="ignore"))]
    if ext == ".docx":
        return [(1, extract_docx_text(file_path))]
    raise ValueError(f"不支持的文件格式: {file_path.name}")

def split_text(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    if chunk_size <= 0:
        raise ValueError("chunk_size 必须大于 0")
    if chunk_overlap < 0:
        raise ValueError("chunk_overlap 不能小于 0")
    if chunk_overlap >= chunk_size:
        raise ValueError("chunk_overlap 必须小于 chunk_size")
    text = text.strip()
    if not text:
        return []
    chunks = []
    start = 0
    length = len(text)
    while start < length:
        end = min(start + chunk_size, length)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= length:
            break
        start = end - chunk_overlap
    return chunks

def collect_chunks(input_dir: Path | List[Path], chunk_size: int, chunk_overlap: int) -> tuple[list[dict], list[dict]]:
    all_chunks = []
    file_stats = []
    
    # 兼容单个目录或文件列表
    if isinstance(input_dir, list):
        files_to_process = sorted(input_dir, key=lambda p: p.name.lower())
    else:
        if not input_dir.exists():
            return [], []
        files_to_process = sorted(input_dir.iterdir(), key=lambda p: p.name.lower())

    for file_path in files_to_process:
        if not file_path.is_file():
            continue
        if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        page_count = 0
        chunk_count = 0
        try:
            pages = load_pages(file_path)
            page_count = len(pages)
            for page_num, raw_page in pages:
                cleaned_page = normalize_text(raw_page)
                page_chunks = split_text(cleaned_page, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
                for idx, chunk_text in enumerate(page_chunks, start=1):
                    chunk_id = f"{file_path.stem}_p{page_num:04d}_c{idx:04d}"
                    all_chunks.append(
                        {
                            "chunk_id": chunk_id,
                            "source_file": file_path.name,
                            "page": page_num,
                            "text": chunk_text,
                        }
                    )
                chunk_count += len(page_chunks)
            file_stats.append(
                {
                    "source_file": file_path.name,
                    "pages": page_count,
                    "chunks": chunk_count,
                }
            )
        except Exception as exc:
            file_stats.append(
                {
                    "source_file": file_path.name,
                    "pages": page_count,
                    "chunks": chunk_count,
                    "error": str(exc),
                }
            )
    return all_chunks, file_stats

def build_hybrid_indices(chunks: list[dict]):
    """为了兼容 module_kb.py 的签名，返回三个占位符，真正的向量构建在 save_outputs 中"""
    return "chroma_vectorizer", "chroma_matrix", "chroma_bm25"

def save_outputs(output_dir: Path, chunks: list[dict], vectorizer, matrix, bm25) -> dict:
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # 1. 依然保存 chunks.jsonl 供其他模块使用（兼容）
    chunks_path = output_dir / "chunks.jsonl"
    with chunks_path.open("w", encoding="utf-8") as f:
        for row in chunks:
            f.write(json.dumps(row, ensure_ascii=False) + "\n")
            
    # 2. 构建 Chroma 向量库 (采用原子交换策略以避免 Windows 文件锁冲突)
    chroma_dir = output_dir / "chroma_db"
    chroma_temp_dir = output_dir / f"chroma_db_tmp_{int(time.time())}"
    
    if chroma_temp_dir.exists():
        shutil.rmtree(chroma_temp_dir)
            
    documents = []
    for c in chunks:
        doc = Document(
            page_content=c["text"],
            metadata={
                "chunk_id": c["chunk_id"],
                "source_file": c["source_file"],
                "page": c["page"]
            }
        )
        documents.append(doc)
        
    embeddings = get_embeddings()
    
    # 先构建到临时目录
    logger.info(f"正在构建向量库到临时目录: {chroma_temp_dir}")
    
    # 使用 PersistentClient 确保更显式的控制
    import chromadb
    from chromadb.config import Settings
    client = chromadb.PersistentClient(path=str(chroma_temp_dir), settings=Settings(anonymized_telemetry=False, is_persistent=True))
    
    vector_store = Chroma.from_documents(
        documents=documents,
        embedding=embeddings,
        client=client,
        persist_directory=str(chroma_temp_dir)
    )
    
    # 显式关闭客户端以释放句柄 (Chroma 0.4.x+)
    try:
        if hasattr(client, "close"):
            client.close()
        elif hasattr(client, "_system") and hasattr(client._system, "stop"):
            client._system.stop()
    except:
        pass
        
    del vector_store
    del client
    import gc
    gc.collect()
    time.sleep(1) # 给系统一点时间完全释放文件句柄

    # 替换旧索引
    logger.info("正在替换旧索引...")
    for i in range(5): 
        try:
            if chroma_dir.exists():
                # Windows 兼容性：先重命名旧目录为 .old，再删除
                old_dir = chroma_dir.with_suffix(".old")
                if old_dir.exists():
                    shutil.rmtree(old_dir, ignore_errors=True)
                chroma_dir.rename(old_dir)
                shutil.rmtree(old_dir, ignore_errors=True)
            
            chroma_temp_dir.rename(chroma_dir)
            logger.info("[success] 索引替换成功")
            break
        except Exception as e:
            if i == 4:
                logger.error(f"替换索引最终失败: {e}")
            else:
                logger.warning(f"替换索引失败 (轮次 {i+1}), 正在重试... {e}")
                time.sleep(2)
                gc.collect()
            
    return {
        "chunks_file": str(chunks_path),
        "chroma_dir": str(chroma_dir),
    }

def load_chunks(chunks_file: Path) -> list[dict]:
    chunks = []
    with chunks_file.open("r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            chunks.append(json.loads(line))
    return chunks

def load_vector_kb(vector_dir: Path):
    """
    返回 (chunks, vectorizer, matrix, bm25) 的结构以兼容现有代码。
    实际上 matrix 位置我们放置初始化的 Chroma 实例。
    """
    chunks_file = vector_dir / "chunks.jsonl"
    chroma_dir = vector_dir / "chroma_db"
    
    if not chunks_file.exists():
        raise FileNotFoundError(f"未找到 chunks 文件: {chunks_file}")
    
    chunks = load_chunks(chunks_file)
    
    # 初始化 Chroma 客户端作为 "matrix"
    embeddings = get_embeddings()
    if chroma_dir.exists():
        chroma_store = Chroma(persist_directory=str(chroma_dir), embedding_function=embeddings)
    else:
        chroma_store = None
        
    # 返回: chunks, vectorizer(dummy), matrix(chroma), bm25(dummy)
    return chunks, "chroma_vectorizer", chroma_store, "chroma_bm25"

def retrieve(question: str, chunks: list[dict], vectorizer, matrix, top_k: int, bm25=None) -> list[dict]:
    """
    参数中的 matrix 实际上是 Chroma 实例。
    """
    chroma_store = matrix
    if not chroma_store:
        logger.warning("Chroma 数据库未初始化，返回空检索结果。")
        return []
        
    query = normalize_text(question)
    enhanced_query = query
    for k, v in QUERY_HINTS.items():
        if k in query:
            enhanced_query = f"{enhanced_query} {v}"
            
    # 执行 Chroma 向量相似度检索 (带分数)
    try:
        try:
            results = chroma_store.similarity_search_with_score(enhanced_query, k=top_k)
        except Exception as e:
            err_msg = str(e).lower()
            # 针对 Windows HNSW 索引加载失败或文件锁冲突进行自动重试/重载
            if "hnsw" in err_msg or "index" in err_msg or "reader" in err_msg:
                logger.warning(f"检测到 Chroma 索引失效 ({e})，尝试重新加载实例...")
                # 尝试从原始持久化目录重新加载
                persist_dir = getattr(chroma_store, "_persist_directory", None)
                if not persist_dir:
                    # 尝试从 settings 中获取
                    try:
                        persist_dir = chroma_store._client._system.settings.persist_directory
                    except:
                        pass
                
                if persist_dir and os.path.exists(persist_dir):
                    embeddings = get_embeddings()
                    # 重新构造实例
                    new_store = Chroma(persist_directory=str(persist_dir), embedding_function=embeddings)
                    results = new_store.similarity_search_with_score(enhanced_query, k=top_k)
                    
                    # 尝试更新全局状态 (如果是在 workflow_engine 环境下)
                    try:
                        from workflow_engine import set_kb_data, KB_CHUNKS, KB_VECTORIZER, KB_BM25, RETRIEVE_FUNC
                        set_kb_data(KB_CHUNKS, KB_VECTORIZER, new_store, RETRIEVE_FUNC, bm25=KB_BM25)
                        logger.info("已成功刷新全局 Chroma 实例。")
                    except:
                        pass
                else:
                    raise e
            else:
                raise e
    except Exception as e:
        logger.error(f"Chroma 检索最终失败: {e}")
        return []
        
    hits = []
    for doc, distance in results:
        # Chroma 默认 L2 distance，转换为相似度得分用于兼容 RRF 和排序
        # 分数转换：score = 1.0 / (1.0 + distance)
        score = round(1.0 / (1.0 + float(distance)), 6)
        
        hits.append({
            "score": score,
            "chunk_id": doc.metadata.get("chunk_id", "unknown"),
            "source_file": doc.metadata.get("source_file", "unknown"),
            "page": doc.metadata.get("page", 1),
            "text": doc.page_content,
            "distance": float(distance) # 供调试
        })
        
    # 按 score 降序排列 (兼容旧逻辑)
    hits = sorted(hits, key=lambda x: x["score"], reverse=True)
    return hits

def build_answer(hits: list[dict]) -> str:
    if not hits:
        return "未检索到有效证据，请尝试更具体的问题。"
    snippets = []
    for hit in hits[:3]:
        text = hit["text"].replace("\n", " ").strip()
        snippets.append(text[:220])
    return " ".join(snippets).strip()

def load_test_questions(questions_file: Path | None) -> list[str]:
    if questions_file is None:
        return DEFAULT_TEST_QUESTIONS
    if not questions_file.exists():
        raise FileNotFoundError(f"测试问题文件不存在: {questions_file}")
    if questions_file.suffix.lower() == ".json":
        data = json.loads(questions_file.read_text(encoding="utf-8"))
        if not isinstance(data, list):
            raise ValueError("JSON 问题文件应为字符串数组")
        questions = [str(x).strip() for x in data if str(x).strip()]
        if not questions:
            raise ValueError("JSON 问题文件为空")
        return questions
    questions = []
    for line in questions_file.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if line:
            questions.append(line)
    if not questions:
        raise ValueError("问题文件为空")
    return questions

def run_rag_test(vector_dir: Path, output_file: Path, top_k: int, questions: list[str]) -> dict:
    chunks, vectorizer, matrix, bm25 = load_vector_kb(vector_dir)
    results = []
    for question in questions:
        hits = retrieve(question, chunks, vectorizer, matrix, top_k=top_k, bm25=bm25)
        citations = [
            {
                "source_file": x["source_file"],
                "page": x["page"],
                "chunk_id": x["chunk_id"],
                "score": x["score"],
            }
            for x in hits
        ]
        results.append(
            {
                "question": question,
                "answer": build_answer(hits),
                "citations": citations,
                "retrieved_chunks": hits,
            }
        )
    report = {
        "vector_dir": str(vector_dir),
        "top_k": top_k,
        "question_count": len(questions),
        "results": results,
    }
    output_file.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return report

def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--mode", choices=["build", "test"], default="build")
    parser.add_argument("--input-dir", default="domain_docs")
    parser.add_argument("--output-dir", default="vector_kb")
    parser.add_argument("--report-file", default="vector_kb_report.json")
    parser.add_argument("--chunk-size", type=int, default=500)
    parser.add_argument("--chunk-overlap", type=int, default=50)
    parser.add_argument("--vector-dir", default="vector_kb")
    parser.add_argument("--test-output", default="rag_test_results.json")
    parser.add_argument("--top-k", type=int, default=5)
    parser.add_argument("--questions-file", default=None)
    args = parser.parse_args()

    if args.mode == "build":
        input_dir = Path(args.input_dir).resolve()
        output_dir = Path(args.output_dir).resolve()
        report_file = Path(args.report_file).resolve()
        if not input_dir.exists():
            raise FileNotFoundError(f"输入目录不存在: {input_dir}")
        chunks, file_stats = collect_chunks(
            input_dir=input_dir,
            chunk_size=args.chunk_size,
            chunk_overlap=args.chunk_overlap,
        )
        if not chunks:
            raise RuntimeError("未生成任何 chunk，无法构建向量库")
        vectorizer, matrix, bm25 = build_hybrid_indices(chunks)
        saved_paths = save_outputs(output_dir, chunks, vectorizer, matrix, bm25)
        report = {
            "input_dir": str(input_dir),
            "output_dir": str(output_dir),
            "chunk_size": args.chunk_size,
            "chunk_overlap": args.chunk_overlap,
            "total_chunks": len(chunks),
            "engine": "ChromaDB + Nomic Embedding",
            "files": file_stats,
            "artifacts": saved_paths,
        }
        report_file.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        print(json.dumps(report, ensure_ascii=False, indent=2))
        return

    vector_dir = Path(args.vector_dir).resolve()
    output_file = Path(args.test_output).resolve()
    questions_file = Path(args.questions_file).resolve() if args.questions_file else None
    questions = load_test_questions(questions_file)
    test_report = run_rag_test(vector_dir=vector_dir, output_file=output_file, top_k=args.top_k, questions=questions)
    print(json.dumps(test_report, ensure_ascii=False, indent=2))

if __name__ == "__main__":
    main()
